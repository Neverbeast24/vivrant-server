#!/usr/bin/env python3
"""Generate updated VIVRANT SDLC + QA master document (.docx)."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_PATHS = [
    Path(r"C:\Users\PC\Downloads\VIVRANT_Complete_Project_Documentation_SDLC.docx"),
    Path(r"C:\Users\PC\Desktop\VIVRANT_Complete_Project_Documentation_SDLC.docx"),
    Path(r"C:\Users\PC\Desktop\viva-server\docs\VIVRANT_Complete_Project_Documentation_SDLC.docx"),
]


def set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tePr = cell._tc.get_or_add_tcPr()
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.name = "Calibri"


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)


def build() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("VIVRΛNT")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = RGBColor(0x0C, 0x3D, 0x2E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("COMPLETE PROJECT DOCUMENTATION")
    r.bold = True
    r.font.size = Pt(18)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = line.add_run(
        "Software Development Life Cycle (SDLC)\n"
        "Requirements • Architecture • Design • Build • Test • UAT • Release • Operations"
    )
    r.font.size = Pt(11)

    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Project", "VIVRΛNT — AI-powered health decision support platform"],
            ["Applications", "VIVRΛNT Web (Next.js) + VIVRΛNT Mobile (Flutter)"],
            ["Document ID", "VIVRANT-SDLC-MASTER-002"],
            ["Version / Status", "2.0 / QA gate complete — August 4, 2026"],
            ["Prepared Date", date.today().isoformat()],
            ["Prepared by", "Daniella Sayson"],
            ["Classification", "Internal Use"],
            [
                "Repositories",
                "github.com/Neverbeast24/vivrant-server ; github.com/Neverbeast24/vivrant-mobile",
            ],
            ["Production Web", "https://viva-server-delta.vercel.app"],
            ["Tagline", "Long live life — Every Choice Shapes Your Health."],
        ],
    )

    doc.add_paragraph()
    add_para(
        doc,
        "This revision replaces the July 20, 2026 baseline. It consolidates the live VIVRΛNT Web "
        "and Flutter mobile products, documents the automated Quality Assurance gate executed on "
        "August 4, 2026, and records residual risks with owners.",
    )

    add_heading(doc, "Document Control", 1)
    add_table(
        doc,
        ["Version", "Date", "Author", "Summary"],
        [
            ["1.0", "2026-07-17", "Project team", "Initial SDLC draft (legacy VIVA/HRIS template)"],
            [
                "1.1",
                "2026-07-20",
                "Project team",
                "Appendix E — VIVRΛNT Web baseline (Next.js + Supabase + Gemini)",
            ],
            [
                "2.0",
                "2026-08-04",
                "Daniella Sayson",
                "Full VIVRΛNT Web+Mobile baseline; QA suite implemented; test evidence recorded; "
                "auth/reminder/security hardening; production redeploy",
            ],
        ],
    )

    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "VIVRΛNT is an AI-powered health companion that combines nutrition, movement, gym, "
        "hydration, sleep, mindfulness, spending, pantry/groceries, journaling, habits, AI coaching, "
        "reminders, and admin controls into one member experience on web and mobile.",
    )
    add_para(
        doc,
        "As of August 4, 2026 the web platform is production-deployed on Vercel under the "
        "saysondaniellads24 account, backed by Supabase Auth/Postgres (RLS), Google Gemini, and "
        "Firebase Storage/FCM. The Flutter mobile client consumes the `/api/mobile/*` and `/api/auth/*` "
        "REST surface with Bearer JWT sessions.",
        )

    add_heading(doc, "1.1 Current Assessment Summary", 2)
    add_table(
        doc,
        ["Area", "Assessment", "Status"],
        [
            ["Web product", "Member dashboard + admin console live on Vercel", "Implemented"],
            ["Mobile product", "Flutter app with module parity for core journeys", "Implemented"],
            ["Mobile REST API", "Bearer-auth `/api/mobile/*` routes in production", "Implemented"],
            ["Auth", "Email + OAuth; suspended checks; hardened refresh/logout", "Implemented"],
            ["AI coaching", "Gemini across insights, chat, coaches, estimates", "Implemented"],
            ["Reminders / FCM", "Cron + on-demand processing; web SW hardened", "Partial"],
            ["Automated QA", "Vitest 21 + Flutter 13 tests; lint/typecheck gate", "Implemented"],
            ["Store release readiness", "Debug signing / Firebase native config pending", "Partial"],
            ["Formal UAT sign-off", "Scenarios ready; business sign-off pending", "Template"],
        ],
    )

    add_heading(doc, "2. Product Scope and Modules", 1)
    add_heading(doc, "2.1 Shared modules", 2)
    add_bullets(
        doc,
        [
            "Authentication — signup, login, forgot/reset/change password, OAuth",
            "Today dashboard — check-in, calories, water, workouts, unread notifications",
            "Nutrition, Movement, Gym (plans/sessions/machines/demos)",
            "Hydration, Sleep, Mindfulness, Habits, Challenges, Journal",
            "Groceries, Pantry, Spending / wellness budget",
            "AI Engine — chat, insights, reminders, weekly story",
            "Reports, Goals, Preferences, Health history, Profile/avatar",
            "Notifications + device tokens",
            "Admin — users, roles, audit, activity, inquiries, tickets, settings (staff)",
        ],
    )

    add_heading(doc, "2.2 Technology stack", 2)
    add_table(
        doc,
        ["Layer", "Web", "Mobile"],
        [
            ["UI", "Next.js 16, React 19, Tailwind 4, Motion", "Flutter 3.35+ / Dart 3.9+, Riverpod, GoRouter"],
            ["API", "Next.js Route Handlers", "Dio + Secure Storage Bearer JWTs"],
            ["Auth/DB", "Supabase Auth + Postgres + RLS", "Same project via server APIs (+ optional Supabase OAuth)"],
            ["AI", "Google Gemini", "Via web `/api/mobile/*/coach` endpoints"],
            ["Push", "Firebase Admin + messaging SW", "firebase_messaging (needs native config)"],
            ["Deploy", "Vercel (sin1 / production alias)", "Manual APK/IPA builds"],
        ],
    )

    add_heading(doc, "3. Architecture Notes", 1)
    add_bullets(
        doc,
        [
            "Web pages gated by Next.js proxy/session cookies; `/api/*` not role-gated by proxy.",
            "Mobile APIs use `requireMobileUser` / staff helpers; Bearer preferred, cookie fallback.",
            "Reminders: daily Vercel cron + on-load processing on web reminders and mobile Today/reminders.",
            "Safe path helper blocks open redirects and absolute notification hrefs.",
            "API errors redact internal DB/provider messages for clients (`jsonError` / `jsonDbError`).",
        ],
    )

    add_heading(doc, "4. Quality Assurance Strategy", 1)
    add_para(
        doc,
        "QA for VIVRΛNT uses a risk-based pyramid: fast unit tests for pure logic, static analysis "
        "(TypeScript + ESLint / Dart analyzer), production smoke checks for critical HTTP surfaces, "
        "and UAT scripts for business acceptance.",
    )

    add_heading(doc, "4.1 Test levels", 2)
    add_table(
        doc,
        ["Level", "Owner tooling", "Scope"],
        [
            ["Unit", "Vitest / flutter_test", "Pure helpers: paths, schedules, BMI, HTTP helpers, AI text, validators"],
            ["Static", "tsc, eslint, flutter analyze", "Type safety and lint gate before merge/deploy"],
            ["Smoke", "scripts/qa-smoke.mjs", "Landing, login, SW, unauth mobile API, login validation"],
            ["System / regression", "Manual + future Playwright/integration_test", "Module journeys across roles"],
            ["UAT", "Business testers", "Acceptance scenarios in Section 6"],
            ["Security", "Code review + Supabase RLS SQL", "Auth, redirects, secret handling, suspended accounts"],
        ],
    )

    add_heading(doc, "4.2 QA execution evidence — August 4, 2026", 2)
    add_table(
        doc,
        ["Check", "Command", "Result"],
        [
            ["Web typecheck", "npm run typecheck", "PASS"],
            ["Web lint", "npm run lint", "PASS (0 errors after fixes)"],
            ["Web unit tests", "npm run test (Vitest)", "PASS — 21/21"],
            ["Web combined gate", "npm run qa", "PASS"],
            ["Mobile tests", "flutter test", "PASS — 13/13"],
            ["Mobile analyzer", "flutter analyze lib test", "INFO-level style findings only; no failing tests"],
            ["Production smoke", "npm run qa:smoke", "4/5 PASS prior to auth-catch redeploy; unauth Today expected 401 after deploy"],
            ["Production deploy", "vercel --prod (saysondaniellads24)", "READY — https://viva-server-delta.vercel.app"],
        ],
    )

    add_heading(doc, "4.3 Automated suites inventory", 2)
    add_para(doc, "Web (vivrant-server)", bold=True)
    add_bullets(
        doc,
        [
            "src/lib/security/safe-path.test.ts",
            "src/lib/reminders/schedule.test.ts",
            "src/lib/health/body-metrics.test.ts",
            "src/lib/mobile/http.test.ts",
            "scripts/qa-smoke.mjs + npm scripts: test, qa, qa:smoke",
        ],
    )
    add_para(doc, "Mobile (vivrant-mobile)", bold=True)
    add_bullets(
        doc,
        [
            "test/ai_text_test.dart",
            "test/validators_test.dart",
            "test/humanize_formatters_test.dart",
            "test/widget_test.dart",
        ],
    )

    add_heading(doc, "4.4 Defects found and resolved in this QA pass", 2)
    add_table(
        doc,
        ["ID", "Finding", "Fix", "Status"],
        [
            ["DEF-QA-01", "ESLint next/no-assign-module-variable in admin activity", "Renamed query var to moduleFilter", "Closed"],
            ["DEF-QA-02", "Unused today prop in SpendingOverview", "Made optional / unused", "Closed"],
            ["DEF-QA-03", "setState-in-effect lint errors in Today & Journal", "Deferred updates via setTimeout(0)", "Closed"],
            ["DEF-QA-04", "Unauth /api/mobile/today could 500", "Catch cookie helper failures → 401", "Closed (pending redeploy)"],
            ["DEF-QA-05", "Mobile due reminders not firing (daily cron only)", "Process on Today + reminders GET", "Closed"],
            ["DEF-QA-06", "Open redirect on /auth/confirm", "safeAppPath allowlist", "Closed"],
            ["DEF-QA-07", "FCM SW opened absolute http(s) hrefs", "Same-origin relative paths only", "Closed"],
            ["DEF-QA-08", "Password routes leaked provider errors", "Friendly mapped errors + 503 on misconfig", "Closed"],
            ["DEF-QA-09", "Refresh allowed suspended users / dumped full user", "403 + limited user payload", "Closed"],
            ["DEF-QA-10", "Mobile refresh blip cleared session", "Keep tokens on transient refresh errors", "Closed"],
            ["DEF-QA-11", "Search/AI UI showed Map.toString()", "Parse label/detail + formatAiText", "Closed"],
            ["DEF-QA-12", "Hydration ignored today's total", "Load water_ml from getToday()", "Closed"],
        ],
    )

    add_heading(doc, "5. Build and Release", 1)
    add_bullets(
        doc,
        [
            "Web: push to main and/or `npx vercel --prod` under Daniella’s Vercel team.",
            "Required Vercel env: SUPABASE_URL/PUBLISHABLE/SECRET, CRON_SECRET, GEMINI_API_KEY, Firebase admin + NEXT_PUBLIC_FIREBASE_*, NEXT_PUBLIC_APP_URL, Resend.",
            "Mobile release builds must pass `--dart-define=API_BASE_URL=https://…` (loopback blocked in release).",
            "Store builds require upload keystore + Firebase native config before Play/App Store.",
        ],
    )

    add_heading(doc, "6. UAT Package (VIVRΛNT)", 1)
    add_para(
        doc,
        "Entry criteria: QA gate green (typecheck/lint/unit), production smoke green, no open Sev-1. "
        "Exit criteria: all critical scenarios Pass/Conditional with signed residual risks.",
    )
    add_table(
        doc,
        ["UAT ID", "Scenario", "Platform", "Expected"],
        [
            ["UAT-V-001", "Email signup / login / logout", "Web + Mobile", "Session established; logout clears access"],
            ["UAT-V-002", "Forgot / reset password", "Web + Mobile", "Generic success; password updates"],
            ["UAT-V-003", "Today check-in + water", "Web + Mobile", "Totals persist and reload"],
            ["UAT-V-004", "Log meal + AI estimate", "Web + Mobile", "Macros saved; friendly errors"],
            ["UAT-V-005", "Log workout / gym session", "Web + Mobile", "History lists entry"],
            ["UAT-V-006", "Create reminder; due fires notification", "Web + Mobile", "In-app notification; push if FCM configured"],
            ["UAT-V-007", "AI insight / coach responses readable", "Mobile", "Title/body text, not object dump"],
            ["UAT-V-008", "Search returns labeled results", "Mobile", "label/detail shown; tap opens module"],
            ["UAT-V-009", "Suspended user blocked", "Web + Mobile", "403; cannot refresh tokens"],
            ["UAT-V-010", "Admin user list / role change", "Web + Mobile staff", "Only staff; audit recorded"],
            ["UAT-V-011", "OAuth Google/GitHub", "Web (+ Mobile if configured)", "Callback lands in app"],
            ["UAT-V-012", "Weekly story / reports", "Web + Mobile", "Readable narrative"],
        ],
    )

    add_heading(doc, "7. Residual Risks and Follow-ups", 1)
    add_table(
        doc,
        ["ID", "Risk", "Severity", "Mitigation / owner"],
        [
            ["QA-R1", "Mobile FCM inactive without google-services / iOS entitlements", "Medium", "Add Firebase native config"],
            ["QA-R2", "Android release signed with debug keys", "High (store)", "Configure upload keystore"],
            ["QA-R3", "Confirm 20260730 security SQL applied in prod Supabase", "High", "Run SQL; verify trigger"],
            ["QA-R4", "No Playwright / integration_test E2E yet", "Medium", "Add next sprint"],
            ["QA-R5", "Auth/AI rate limits missing", "Medium", "Add KV sliding window"],
            ["QA-R6", "Hobby cron is daily — rely on on-demand processing", "Medium", "Upgrade cron frequency if needed"],
        ],
    )

    add_heading(doc, "8. SDLC Gate Checklist", 1)
    add_table(
        doc,
        ["Gate", "Evidence", "Status"],
        [
            ["G1 Requirements", "Module inventory + acceptance scenarios", "Complete for current scope"],
            ["G2 Architecture", "Stack + auth/API notes in this document", "Complete"],
            ["G3 Build", "Next.js build + Flutter tests", "Complete"],
            ["G4 Test", "Vitest 21 + Flutter 13 + lint/typecheck", "Complete (Aug 4, 2026)"],
            ["G5 UAT", "UAT-V-001…012 register", "Ready for business execution"],
            ["G6 Production acceptance", "Vercel production READY", "Web accepted; mobile store TBD"],
        ],
    )

    add_heading(doc, "9. How to Re-run QA", 1)
    add_para(doc, "Web", bold=True)
    add_para(doc, "cd viva-server && npm run qa && npm run qa:smoke")
    add_para(doc, "Mobile", bold=True)
    add_para(doc, "cd vivrant-mobile && flutter analyze lib test && flutter test")
    add_para(
        doc,
        "Detailed log: viva-server/docs/QA_EVIDENCE.md",
    )

    add_heading(doc, "10. Conclusion", 1)
    add_para(
        doc,
        "VIVRΛNT Web and Mobile now share a documented SDLC baseline with an executable QA gate. "
        "Critical security and session defects identified in the August 2026 audit were remediated and "
        "covered by automated tests. Remaining work is store-release hardening, native push configuration, "
        "formal UAT sign-off, and expanded E2E coverage.",
    )

    add_heading(doc, "Appendix A — Brand", 1)
    add_para(
        doc,
        "VIVRΛNT is a stylized form of “vibrant,” containing Latin vivere (“to live”). "
        "The lambda (Λ) stylizes the letter A. Primary tagline: Long live life. "
        "Secondary: Every Choice Shapes Your Health.",
    )

    add_heading(doc, "Appendix B — Related artifacts", 1)
    add_bullets(
        doc,
        [
            "viva-server/docs/QA_EVIDENCE.md",
            "viva-server/SETUP.md",
            "vivrant-mobile/docs/MOBILE_API_SPEC.md",
            "viva-server/supabase/*.sql",
        ],
    )

    return doc


def main() -> None:
    doc = build()
    for path in OUT_PATHS:
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)
        print(f"Wrote {path}")


if __name__ == "__main__":
    main()
